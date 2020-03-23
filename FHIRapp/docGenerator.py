import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


### This is not a runnable app, but merely an auxiliary file I created for better grouping of the code


# takes a patient's ordered information and writes it into the given document neatly
def writeList(title, mainBody, dictionary):
    mainBody.add_run(title + ": \n").bold = True
    count = 0
    for item in dictionary:
        count += 1
        mainBody.add_run(str(count) + ". " + item + ": " + dictionary[item] + "\n")
        if count == len(dictionary):
            mainBody.add_run("\n")


# takes a dictionary of all the patients informations, from the API, and then creates a document
def generateHealthDoc(patientDictionary):
    doc = docx.Document()
    doc.add_heading("Patient Blood Report", 0)
    addressPara = doc.add_paragraph(patientDictionary['Address'])
    addressPara.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(patientDictionary['Name'])
    mainBody = doc.add_paragraph()
    writeList("Your full blood count", mainBody, patientDictionary['Full-Blood Count Attributes'])
    writeList("Blood Pressure readings", mainBody, patientDictionary['Blood Pressure Values'])
    writeList("Other blood readings", mainBody, patientDictionary['Other blood readings'])
    doc.save(patientDictionary['UID'] + ".docx")
